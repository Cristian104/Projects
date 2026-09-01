import os
import io
import zipfile
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from fpdf import FPDF
from typing import Dict, Any

def create_cover_letter_docx(job_data: Dict[str, Any], cover_text: str) -> bytes:
    """Generate professional DOCX cover letter with full UTF-8 Polish support."""
    doc = Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Header Name
    p_name = doc.add_paragraph()
    run_name = p_name.add_run("Kamila Drewniak")
    run_name.bold = True
    run_name.font.size = Pt(18)
    run_name.font.color.rgb = RGBColor(16, 185, 129) # Emerald
    
    # Contact Info
    p_contact = doc.add_paragraph()
    run_contact = p_contact.add_run("Email: kamila_dre@interia.pl | Phone: +48 888 788 085 | Location: Częstochowa, Polska")
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_paragraph() # Spacer
    
    # Recipient & Job Info
    company = job_data.get("company", "Dział Rekrutacji")
    title = job_data.get("title", "Stanowisko")
    
    p_rec = doc.add_paragraph()
    r_rec = p_rec.add_run(f"Do: Dział Rekrutacji / HR\n{company}\nDotyczy aplikacji na stanowisko: {title}\n")
    r_rec.font.size = Pt(11)
    r_rec.bold = True
    
    doc.add_paragraph() # Spacer
    
    # Body Paragraphs
    paragraphs = cover_text.strip().split("\n\n")
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            r = p.add_run(para.strip())
            r.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.25
            p.paragraph_format.space_after = Pt(8)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def normalize_polish_text(text: str) -> str:
    """Safely map Polish diacritics for standard FPDF Helvetica engine."""
    replacements = {
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n', 'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N', 'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

class PDFCoverLetter(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 16)
        self.set_text_color(16, 185, 129)
        self.cell(0, 8, 'Kamila Drewniak', ln=True)
        
        self.set_font('Helvetica', '', 9)
        self.set_text_color(100, 116, 139)
        self.cell(0, 5, 'Email: kamila_dre@interia.pl | Phone: +48 888 788 085 | Location: Czestochowa, Polska', ln=True)
        self.ln(6)
        
        self.set_draw_color(226, 232, 240)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(6)

def create_cover_letter_pdf(job_data: Dict[str, Any], cover_text: str) -> bytes:
    """Generate clean PDF cover letter using FPDF2."""
    pdf = PDFCoverLetter()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    company = normalize_polish_text(job_data.get("company", "Dzial Rekrutacji"))
    title = normalize_polish_text(job_data.get("title", "Stanowisko"))
    safe_text = normalize_polish_text(cover_text)
    
    pdf.set_font('Helvetica', 'B', 11)
    pdf.set_text_color(15, 23, 42)
    pdf.multi_cell(0, 6, f"Do: Dzial Rekrutacji / HR\nFirma: {company}\nDotyczy: {title}\n")
    pdf.ln(5)
    
    pdf.set_font('Helvetica', '', 10.5)
    pdf.multi_cell(0, 6, safe_text)
    
    return bytes(pdf.output())

def build_application_zip(job_data: Dict[str, Any], cover_text: str, include_cv: bool = True, cv_path: str = "CV_KAMILA_DREWNIAK.pdf") -> bytes:
    """Bundle DOCX cover letter, PDF cover letter, and candidate CV into a ZIP archive."""
    company_clean = "".join(c for c in job_data.get("company", "Company") if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    title_clean = "".join(c for c in job_data.get("title", "Job") if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    
    docx_bytes = create_cover_letter_docx(job_data, cover_text)
    pdf_bytes = create_cover_letter_pdf(job_data, cover_text)
    
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"CoverLetter_{company_clean}_{title_clean}.docx", docx_bytes)
        zf.writestr(f"CoverLetter_{company_clean}_{title_clean}.pdf", pdf_bytes)
        
        if include_cv and os.path.exists(cv_path):
            with open(cv_path, "rb") as f_cv:
                zf.writestr("CV_Kamila_Drewniak.pdf", f_cv.read())
                
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
